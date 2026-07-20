import streamlit as st
import os
os.environ["HTTP_PROXY"] = ""
os.environ["HTTPS_PROXY"] = ""
os.environ["http_proxy"] = ""
os.environ["https_proxy"] = ""
os.environ["NO_PROXY"] = "*"
os.environ["no_proxy"] = "*"
import sys
import uuid
import json
import asyncio
import time
import httpx
from datetime import datetime, timezone
from pathlib import Path
import plotly.graph_objects as go
import numpy as np

# Ensure backend folder is in sys.path
BASE_DIR = Path(__file__).resolve().parent
backend_path = BASE_DIR / "backend"
if str(backend_path) not in sys.path:
    sys.path.insert(0, str(backend_path))

# In standalone mode, we can import backend functions directly.
try:
    from app.config import settings
    from app.database import init_db, get_db_context
    from app.models.db_models import Project, GeometryFeatures, Job, AnalysisResult, Material, ChatMessage, JobStatus, AnalysisType
    from app.workers.seed import seed_materials
    from app.cad_import.parser import parse_cad_file
    from app.geometry_analysis.feature_extractor import extract_features
    from app.geometry_analysis.suggestion_engine import suggest_analyses
    from app.workers.inline_runner import run_analysis_inline
    from app.chatbot.gemini_client import stream_chat_response
    STANDALONE_AVAILABLE = True
except Exception as e:
    STANDALONE_AVAILABLE = False
    STANDALONE_ERROR = e

# Set up page configurations
st.set_page_config(
    page_title="CAD-CAE Analyzer — Streamlit Dashboard",
    page_icon="⚙️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for modern design aesthetics matching the React dashboard
if "theme" not in st.session_state:
    st.session_state.theme = "Dark"

def apply_theme_css():
    if st.session_state.theme == "Light":
        st.markdown("""
        <style>
            .stApp {
                background-color: #f8f9fa !important;
                color: #212529 !important;
            }
            h1, h2, h3, h4, h5, h6 {
                color: #1976d2 !important;
                font-family: 'Inter', system-ui, -apple-system, sans-serif !important;
            }
            p, span, label, li, .stMarkdown p {
                color: #495057 !important;
            }
            /* Buttons */
            .stButton>button {
                background: linear-gradient(135deg, #1565c0, #1976d2) !important;
                color: white !important;
                font-weight: 600 !important;
                border-radius: 8px !important;
                border: none !important;
                padding: 0.5rem 1.5rem !important;
                box-shadow: 0 4px 12px rgba(25, 118, 210, 0.2) !important;
            }
            /* Sidebar */
            [data-testid="stSidebar"] {
                background-color: #e9ecef !important;
            }
            [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 {
                color: #1565c0 !important;
            }
            [data-testid="stSidebar"] p, [data-testid="stSidebar"] span, [data-testid="stSidebar"] label {
                color: #495057 !important;
            }
            /* Cards styling */
            .metric-card {
                background: #ffffff !important;
                border: 1px solid #dee2e6 !important;
                border-radius: 12px;
                padding: 20px;
                margin-bottom: 15px;
                box-shadow: 0 4px 12px rgba(0,0,0,0.05) !important;
            }
            .metric-title {
                color: #6c757d !important;
            }
            .metric-value {
                color: #212529 !important;
            }
            .suggestion-card {
                background: #f1f3f5 !important;
                border: 1px solid #ced4da !important;
                border-radius: 10px;
                padding: 16px;
                margin-bottom: 12px;
            }
            .suggestion-title {
                color: #1976d2 !important;
            }
            /* Form / inputs */
            div[data-baseweb="select"] > div {
                background-color: #ffffff !important;
                color: #212529 !important;
                border-color: #ced4da !important;
            }
            div[data-baseweb="input"] input {
                background-color: #ffffff !important;
                color: #212529 !important;
            }
        </style>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <style>
            .stApp {
                background-color: #060d1a !important;
                color: #e8f0fe !important;
            }
            h1, h2, h3, h4, h5, h6 {
                color: #00e5ff !important;
                font-family: 'Inter', system-ui, -apple-system, sans-serif !important;
            }
            p, span, label, li, .stMarkdown p {
                color: #8ba3c7 !important;
            }
            /* Buttons */
            .stButton>button {
                background: linear-gradient(135deg, #2979ff, #00e5ff) !important;
                color: white !important;
                font-weight: 600 !important;
                border-radius: 8px !important;
                border: none !important;
                padding: 0.5rem 1.5rem !important;
                box-shadow: 0 4px 16px rgba(0, 229, 255, 0.2) !important;
                transition: all 0.2s !important;
            }
            .stButton>button:hover {
                transform: translateY(-1px) !important;
                box-shadow: 0 6px 20px rgba(0, 229, 255, 0.4) !important;
            }
            /* Cards styling */
            .metric-card {
                background: rgba(17, 34, 64, 0.85) !important;
                border: 1px solid rgba(41, 121, 255, 0.2) !important;
                border-radius: 12px;
                padding: 20px;
                margin-bottom: 15px;
                box-shadow: 0 8px 32px rgba(0,0,0,0.4) !important;
                backdrop-filter: blur(12px);
            }
            .metric-title {
                color: #4a6080 !important;
            }
            .metric-value {
                color: #e8f0fe !important;
            }
            .suggestion-card {
                background: rgba(23, 42, 74, 0.7) !important;
                border: 1px solid rgba(0, 229, 255, 0.25) !important;
                border-radius: 10px;
                padding: 16px;
                margin-bottom: 12px;
            }
            .suggestion-title {
                color: #00e5ff !important;
            }
        </style>
        """, unsafe_allow_html=True)

apply_theme_css()

# ── Helper for Running Async Coroutines in Streamlit ─────────────────────────
def run_async(coro):
    try:
        loop = asyncio.get_running_loop()
    except RuntimeError:
        loop = None
    if loop and loop.is_running():
        import concurrent.futures
        with concurrent.futures.ThreadPoolExecutor() as executor:
            future = executor.submit(asyncio.run, coro)
            return future.result()
    else:
        return asyncio.run(coro)

# ── Initialize DB & Materials in Standalone mode ─────────────────────────────
@st.cache_resource
def setup_database():
    if STANDALONE_AVAILABLE:
        run_async(init_db())
        run_async(seed_materials())
        return True
    return False

db_ready = setup_database()

# ── Session State Setup ───────────────────────────────────────────────────────
st.session_state.op_mode = "Standalone"

if st.session_state.get("user") != "default_user" or st.session_state.get("token") is not None:
    st.session_state.user = "default_user"
    st.session_state.user_id = "default_user_id"
    st.session_state.token = None

if "active_project_id" not in st.session_state:
    st.session_state.active_project_id = None
if "active_job_id" not in st.session_state:
    st.session_state.active_job_id = None
if "chat_messages" not in st.session_state:
    st.session_state.chat_messages = []
if "active_tab" not in st.session_state:
    st.session_state.active_tab = "📐 Geometry Analysis"

# ── API Client Implementation ────────────────────────────────────────────────
def api_request(method, endpoint, **kwargs):
    url = f"{st.session_state.api_url}{endpoint}"
    try:
        # Attach Bearer Token to headers if it exists in session state
        headers = kwargs.get("headers", {})
        if "token" in st.session_state and st.session_state.token:
            headers["Authorization"] = f"Bearer {st.session_state.token}"
        kwargs["headers"] = headers
        
        response = httpx.request(method, url, timeout=30.0, **kwargs)
        if response.status_code >= 400:
            st.error(f"API Error ({response.status_code}): {response.text}")
            return None
        return response.json()
    except Exception as e:
        st.error(f"Failed to connect to API at {url}: {e}")
        if STANDALONE_AVAILABLE:
            st.info("💡 **Tip**: The API backend appears to be offline. You can switch to **Standalone** mode in the sidebar to run the application completely locally.")
        return None


def api_chat_stream(api_url, project_id, job_id, message):
    url = f"{api_url}/chat"
    payload = {
        "project_id": project_id,
        "job_id": job_id,
        "message": message,
        "stream": True
    }
    try:
        with httpx.stream("POST", url, json=payload, timeout=60.0) as r:
            if r.status_code >= 400:
                yield f"API Error ({r.status_code}): {r.read().decode()}"
                return
            for line in r.iter_lines():
                if line.startswith("data: "):
                    data_str = line[6:]
                    try:
                        data_json = json.loads(data_str)
                        if data_json.get("done"):
                            break
                        if "token" in data_json:
                            yield data_json["token"]
                        elif "error" in data_json:
                            yield f"\n[Error: {data_json['error']}]\n"
                    except json.JSONDecodeError:
                        pass
    except Exception as e:
        yield f"Failed to connect to chat API: {e}"


def local_upload_cad(file_bytes, filename, project_name, user_id=None):
    if not STANDALONE_AVAILABLE:
        st.error("Standalone imports failed. Check standard library packages.")
        return None

    # Generate project ID
    project_id = str(uuid.uuid4())
    
    # Save the file
    upload_dir = settings.upload_dir / project_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    saved_path = upload_dir / filename
    saved_path.write_bytes(file_bytes)
    
    # Parse geometry and extract features
    async def _parse_and_persist():
        geom_data = await parse_cad_file(saved_path)
        features = extract_features(geom_data)
        suggestions = suggest_analyses(features, None)
        
        name = project_name or Path(filename).stem
        suffix = Path(filename).suffix.lower().lstrip(".")
        
        async with get_db_context() as db:
            project = Project(
                id=project_id,
                name=name,
                cad_filename=filename,
                cad_format=suffix,
                cad_file_path=str(saved_path),
                cad_file_size=len(file_bytes),
                user_id=user_id,
            )
            db.add(project)
            
            com = features.center_of_mass
            geo_db = GeometryFeatures(
                project_id=project_id,
                volume=features.volume,
                surface_area=features.surface_area,
                bbox_x=features.bbox[0],
                bbox_y=features.bbox[1],
                bbox_z=features.bbox[2],
                center_of_mass=f"{com[0]:.2f},{com[1]:.2f},{com[2]:.2f}",
                slenderness_ratio=features.slenderness_ratio,
                surface_volume_ratio=features.surface_volume_ratio,
                is_thin_walled=features.is_thin_walled,
                has_holes=features.has_holes,
                has_fillets=features.has_fillets,
                has_internal_cavity=features.has_internal_cavity,
                has_symmetry=features.has_symmetry,
                is_sheet_metal=features.is_sheet_metal,
                raw_features=features.raw,
                suggestions=[
                    {
                        "analysis_type": s.analysis_type.value,
                        "title": s.title,
                        "rationale": s.rationale,
                        "confidence": s.confidence,
                        "icon": s.icon,
                        "category": s.category,
                        "priority": s.priority,
                    }
                    for s in suggestions
                ],
            )
            db.add(geo_db)
            
    run_async(_parse_and_persist())
    return project_id


# ── Standalone Mode Database Wrappers ────────────────────────────────────────
def get_standalone_projects(user_id=None):
    async def _fetch():
        async with get_db_context() as db:
            from sqlalchemy import select, or_
            stmt = select(Project)
            if user_id:
                stmt = stmt.where(or_(Project.user_id == user_id, Project.user_id.is_(None)))
            else:
                stmt = stmt.where(Project.user_id.is_(None))
            stmt = stmt.order_by(Project.created_at.desc())
            rows = (await db.execute(stmt)).scalars().all()
            return [{"id": p.id, "name": p.name, "cad_filename": p.cad_filename, "cad_format": p.cad_format} for p in rows]
    return run_async(_fetch()) if STANDALONE_AVAILABLE else []

def get_standalone_project(project_id, user_id=None):
    async def _fetch():
        async with get_db_context() as db:
            from sqlalchemy import select, or_
            stmt = select(Project).where(Project.id == project_id)
            if user_id:
                stmt = stmt.where(or_(Project.user_id == user_id, Project.user_id.is_(None)))
            p = (await db.execute(stmt)).scalar_one_or_none()
            if not p:
                return None
            g = (await db.execute(select(GeometryFeatures).where(GeometryFeatures.project_id == project_id))).scalar_one_or_none()
            
            geo_dict = {}
            if g:
                com = [float(x) for x in g.center_of_mass.split(",")] if g.center_of_mass else [0, 0, 0]
                geo_dict = {
                    "volume": g.volume,
                    "surface_area": g.surface_area,
                    "bbox_x": g.bbox_x, "bbox_y": g.bbox_y, "bbox_z": g.bbox_z,
                    "center_of_mass": com,
                    "slenderness_ratio": g.slenderness_ratio,
                    "surface_volume_ratio": g.surface_volume_ratio,
                    "is_thin_walled": g.is_thin_walled,
                    "has_holes": g.has_holes,
                    "has_fillets": g.has_fillets,
                    "has_internal_cavity": g.has_internal_cavity,
                    "has_symmetry": g.has_symmetry,
                    "is_sheet_metal": g.is_sheet_metal,
                    "suggestions": g.suggestions or [],
                }
            return {
                "id": p.id, "name": p.name, "cad_filename": p.cad_filename,
                "cad_format": p.cad_format, "cad_file_path": p.cad_file_path,
                "geometry": geo_dict
            }
    return run_async(_fetch()) if STANDALONE_AVAILABLE else None

def get_standalone_materials():
    async def _fetch():
        async with get_db_context() as db:
            from sqlalchemy import select
            mats = (await db.execute(select(Material).order_by(Material.name))).scalars().all()
            return [
                {
                    "id": m.id, "name": m.name, "category": m.category,
                    "youngs_modulus": m.youngs_modulus, "poissons_ratio": m.poissons_ratio,
                    "density": m.density, "yield_strength": m.yield_strength,
                    "ultimate_strength": m.ultimate_strength,
                    "thermal_conductivity": m.thermal_conductivity,
                    "specific_heat": m.specific_heat, "thermal_expansion": m.thermal_expansion
                }
                for m in mats
            ]
    return run_async(_fetch()) if STANDALONE_AVAILABLE else []

def get_standalone_jobs(project_id):
    async def _fetch():
        async with get_db_context() as db:
            from sqlalchemy import select
            rows = (await db.execute(select(Job).where(Job.project_id == project_id).order_by(Job.created_at.desc()))).scalars().all()
            return [{"id": j.id, "analysis_types": j.analysis_types, "analysis_type": j.analysis_types[0] if j.analysis_types else "static_structural", "status": j.status, "progress_percent": j.progress_percent} for j in rows]
    return run_async(_fetch()) if STANDALONE_AVAILABLE else []

def get_standalone_job(job_id):
    async def _fetch():
        async with get_db_context() as db:
            from sqlalchemy import select
            j = (await db.execute(select(Job).where(Job.id == job_id))).scalar_one_or_none()
            if not j:
                return None
            res = (await db.execute(select(AnalysisResult).where(AnalysisResult.job_id == job_id))).scalar_one_or_none()
            res_dict = {}
            if res:
                res_dict = {
                    "max_stress": res.max_stress,
                    "min_stress": res.min_stress,
                    "max_displacement": res.max_displacement,
                    "max_temperature": res.max_temperature,
                    "min_safety_factor": res.min_safety_factor,
                    "natural_frequencies": res.natural_frequencies,
                    "buckling_factors": res.buckling_factors,
                    "fatigue_life_cycles": res.fatigue_life_cycles,
                    "result_data": res.result_data or {},
                }
            return {
                "id": j.id, "analysis_types": j.analysis_types, "analysis_type": j.analysis_types[0] if j.analysis_types else "static_structural", "status": j.status,
                "progress_percent": j.progress_percent, "progress_message": j.progress_message,
                "error_message": j.error_message, "results": res_dict
            }
    return run_async(_fetch()) if STANDALONE_AVAILABLE else None


# ── Authentication Helper Functions ──────────────────────────────────────────
import hashlib

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(password: str, hashed: str) -> bool:
    return hash_password(password) == hashed

def check_standalone_login(username, password):
    async def _check():
        async with get_db_context() as db:
            from sqlalchemy import select
            from app.models.db_models import User
            user = (await db.execute(select(User).where(User.username == username))).scalar_one_or_none()
            if user and verify_password(password, user.hashed_password):
                return user.id
            return None
    return run_async(_check()) if STANDALONE_AVAILABLE else None

def create_standalone_user(username, password):
    async def _create():
        async with get_db_context() as db:
            from sqlalchemy import select
            from app.models.db_models import User
            existing = (await db.execute(select(User).where(User.username == username))).scalar_one_or_none()
            if existing:
                return "exists"
            new_user = User(username=username, hashed_password=hash_password(password))
            db.add(new_user)
            await db.commit()
            return new_user.id
    return run_async(_create()) if STANDALONE_AVAILABLE else None

def check_api_login(username, password):
    url = f"{st.session_state.api_url}/auth/login"
    try:
        response = httpx.post(url, json={"username": username, "password": password}, timeout=10.0)
        if response.status_code == 200:
            data = response.json()
            return data.get("token"), data.get("user_id")
        else:
            st.error(f"Login failed: {response.text}")
            return None, None
    except Exception as e:
        st.error(f"Failed to connect to API at {url}: {e}")
        return None, None

def create_api_user(username, password):
    url = f"{st.session_state.api_url}/auth/signup"
    try:
        response = httpx.post(url, json={"username": username, "password": password}, timeout=10.0)
        if response.status_code == 201:
            data = response.json()
            return data.get("user_id") or "success"
        elif response.status_code == 400:
            return "exists"
        else:
            st.error(f"Sign up failed: {response.text}")
            return None
    except Exception as e:
        st.error(f"Failed to connect to API at {url}: {e}")
        return None
def get_plotly_mesh(vertices, faces, analysis_type=None, result_data=None, field_type="stress", slice_axis=None, slice_val=None):
    """
    Renders CAD geometry or FEA solver result fields onto the 3D model using Plotly.
    Supports dynamic slicing plane cuts and 3D velocity vectors (Cone) for CFD.
    """
    if vertices is None or len(vertices) == 0:
        return None

    verts = np.array(vertices)
    x, y, z = verts[:, 0], verts[:, 1], verts[:, 2]
    faces_arr = np.array(faces)

    # Perform cutaway if slicing plane is active
    if slice_axis and slice_axis != "None" and slice_val is not None:
        axis_idx = 0 if slice_axis == "X-Axis" else (1 if slice_axis == "Y-Axis" else 2)
        keep_faces = []
        for face in faces_arr:
            if np.all(verts[face, axis_idx] <= slice_val):
                keep_faces.append(face)
        if len(keep_faces) > 0:
            faces_arr = np.array(keep_faces)

    i, j, k = faces_arr[:, 0], faces_arr[:, 1], faces_arr[:, 2]

    title = "CAD Geometry Mesh"
    color_intensity = None
    colorscale = "Viridis"
    colorbar_title = ""

    # Generate synthetic results mapped to geometry vertices if in mock/emulation mode
    if result_data:
        num_verts = len(verts)
        max_val = 0.0
        min_val = 0.0

        # Create realistic color contours depending on simulation type
        if analysis_type in ("static_structural", "nonlinear", "fatigue"):
            # Mock stress/displacement gradient (higher near the fixed support (x_min), lower at tips)
            min_x, max_x = np.min(x), np.max(x)
            span_x = max(max_x - min_x, 1.0)
            
            if field_type == "displacement":
                # Displacement is maximum at the tip (x_max)
                disp_max = result_data.get("max_displacement_mm", result_data.get("max_displacement", 0.05))
                color_intensity = disp_max * ((x - min_x) / span_x)**2
                colorbar_title = "Displacement (mm)"
                colorscale = "Plasma"
                title = f"3D Displacement Field (Max: {disp_max:.4f} mm)"
            elif field_type == "safety_factor":
                sf_min = result_data.get("min_safety_factor", 1.5)
                # Safety factor is lowest where stress is highest (at x_min)
                color_intensity = sf_min + (10 - sf_min) * ((x - min_x) / span_x)
                colorbar_title = "Safety Factor"
                colorscale = "Portland"
                title = f"Safety Factor Contour (Min: {sf_min:.2f})"
            else:  # stress
                stress_max = result_data.get("max_stress_mpa", result_data.get("max_stress", 150.0))
                # Stress is maximum at the fixed support
                color_intensity = stress_max * ((max_x - x) / span_x) + np.random.uniform(0, 10, num_verts)
                color_intensity = np.clip(color_intensity, 0, stress_max)
                colorbar_title = "Stress (MPa)"
                colorscale = "Jet"
                title = f"3D von Mises Stress (Max: {stress_max:.1f} MPa)"

        elif "thermal" in str(analysis_type):
            temp_max = result_data.get("max_temperature", 100.0)
            temp_min = result_data.get("min_temperature", result_data.get("min_temperature_c", 20.0))
            # Temperature gradient from bottom to top (y_min to y_max)
            min_y, max_y = np.min(y), np.max(y)
            span_y = max(max_y - min_y, 1.0)
            color_intensity = temp_min + (temp_max - temp_min) * ((y - min_y) / span_y)
            colorbar_title = "Temperature (°C)"
            colorscale = "Hot"
            title = f"Steady-State Thermal Distribution (Max: {temp_max:.1f} °C)"

        elif analysis_type in ("cfd_internal", "cfd_external"):
            vel_max = result_data.get("max_velocity", 5.0)
            # Velocity contours (higher near center of bbox)
            centroid = np.mean(verts, axis=0)
            dist_to_center = np.linalg.norm(verts - centroid, axis=1)
            max_dist = max(np.max(dist_to_center), 1.0)
            color_intensity = vel_max * (1.0 - (dist_to_center / max_dist)**2)
            colorbar_title = "Velocity (m/s)"
            colorscale = "Viridis"
            title = f"CFD Flow Velocity field (Max: {vel_max:.2f} m/s)"

        elif analysis_type in ("modal", "buckling"):
            # Sinusoidal mode shapes
            min_x, max_x = np.min(x), np.max(x)
            span_x = max(max_x - min_x, 1.0)
            color_intensity = np.sin((x - min_x) / span_x * np.pi * 2)
            colorbar_title = "Relative Amplitude"
            colorscale = "RdBu"
            title = "Modal Deformation Mode Shape"

    # Define Plotly Mesh3d Trace
    mesh_trace = go.Mesh3d(
        x=x, y=y, z=z,
        i=i, j=j, k=k,
        intensity=color_intensity,
        colorscale=colorscale,
        colorbar=dict(title=colorbar_title, thickness=20) if color_intensity is not None else None,
        color="#2979ff" if color_intensity is None else None,
        flatshading=True,
        lighting=dict(ambient=0.5, diffuse=0.8, specular=0.2, roughness=0.5),
        lightposition=dict(x=100, y=100, z=100)
    )

    data_traces = [mesh_trace]

    # Add 3D Velocity vectors (Cone) for CFD results
    if result_data and analysis_type in ("cfd_internal", "cfd_external"):
        vel_max = result_data.get("max_velocity", 5.0)
        stride = max(1, len(verts) // 40)
        arrow_x = x[::stride]
        arrow_y = y[::stride]
        arrow_z = z[::stride]
        centroid = np.mean(verts, axis=0)
        dists = np.linalg.norm(verts[::stride] - centroid, axis=1)
        max_dist = max(np.max(dists), 1.0)
        v_mags = vel_max * (1.0 - (dists / max_dist)**2)
        
        u = v_mags * 0.95
        v = (arrow_y - centroid[1]) * 0.05 * v_mags
        w = (arrow_z - centroid[2]) * 0.05 * v_mags
        
        cone_trace = go.Cone(
            x=arrow_x, y=arrow_y, z=arrow_z,
            u=u, v=v, w=w,
            sizemode="scaled",
            sizeref=2.0,
            colorscale="Viridis",
            showscale=False,
            name="Velocity Vectors"
        )
        data_traces.append(cone_trace)

    # Add semi-transparent cutting plane visual
    if slice_axis and slice_axis != "None" and slice_val is not None:
        min_x, max_x = np.min(x), np.max(x)
        min_y, max_y = np.min(y), np.max(y)
        min_z, max_z = np.min(z), np.max(z)
        if slice_axis == "X-Axis":
            plane_x = [[slice_val, slice_val], [slice_val, slice_val]]
            plane_y = [[min_y, max_y], [min_y, max_y]]
            plane_z = [[min_z, min_z], [max_z, max_z]]
        elif slice_axis == "Y-Axis":
            plane_x = [[min_x, max_x], [min_x, max_x]]
            plane_y = [[slice_val, slice_val], [slice_val, slice_val]]
            plane_z = [[min_z, min_z], [max_z, max_z]]
        else:
            plane_x = [[min_x, max_x], [min_x, max_x]]
            plane_y = [[min_y, min_y], [max_y, max_y]]
            plane_z = [[slice_val, slice_val], [slice_val, slice_val]]
        
        plane_trace = go.Surface(
            x=plane_x, y=plane_y, z=plane_z,
            colorscale=[[0, 'rgba(255, 23, 68, 0.45)'], [1, 'rgba(255, 23, 68, 0.45)']],
            showscale=False,
            name="Cutting Plane"
        )
        data_traces.append(plane_trace)

    fig = go.Figure(data=data_traces)
    is_light = st.session_state.get("theme") == "Light"
    paper_bg = "#ffffff" if is_light else "#0d1b2e"
    title_color = "#1976d2" if is_light else "#00e5ff"
    fig.update_layout(
        title=dict(text=title, font=dict(color=title_color)),
        paper_bgcolor=paper_bg,
        plot_bgcolor=paper_bg,
        margin=dict(l=0, r=0, t=40, b=0),
        scene=dict(
            xaxis=dict(visible=False),
            yaxis=dict(visible=False),
            zaxis=dict(visible=False),
            bgcolor=paper_bg
        ),
        height=500
    )
    return fig


def generate_mock_pdf(project_name, analysis_type, results):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        import io
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            name='TitleStyle',
            parent=styles['Heading1'],
            textColor=colors.HexColor('#002d62'),
            spaceAfter=20
        )
        
        story.append(Paragraph(f"CAD-CAE Simulation Analysis Report", title_style))
        story.append(Paragraph(f"<b>Project Name:</b> {project_name}", styles['Normal']))
        story.append(Paragraph(f"<b>Simulation Type:</b> {analysis_type.replace('_', ' ').title()}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("<b>Simulation Scalar Summary:</b>", styles['Heading3']))
        story.append(Spacer(1, 6))
        
        data = [["Metric", "Value", "Unit"]]
        if results.get("max_stress"):
            data.append(["Max von Mises Stress", f"{results.get('max_stress'):,.2f}", "MPa"])
        if results.get("max_displacement"):
            data.append(["Max Displacement", f"{results.get('max_displacement'):,.4f}", "mm"])
        if results.get("min_safety_factor"):
            data.append(["Min Safety Factor", f"{results.get('min_safety_factor'):,.2f}", "-"])
        if results.get("max_temperature"):
            data.append(["Max Temperature", f"{results.get('max_temperature'):,.1f}", "°C"])
            
        t = Table(data, colWidths=[200, 100, 100])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#2979ff')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 8),
            ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#f5f7fa')),
            ('GRID', (0,0), (-1,-1), 1, colors.HexColor('#d0d5dd')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')])
        ]))
        story.append(t)
        story.append(Spacer(1, 20))
        story.append(Paragraph("Report generated automatically via CAD-CAE dashboard.", styles['Italic']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        return f"CAD-CAE Simulation Analysis Report\nProject: {project_name}\nType: {analysis_type}\nError generating ReportLab PDF: {e}".encode()


def generate_mock_docx(project_name, analysis_type, results):
    try:
        from docx import Document
        import io
        
        doc = Document()
        doc.add_heading("CAD-CAE Simulation Analysis Report", level=0)
        
        doc.add_paragraph(f"Project Name: {project_name}")
        doc.add_paragraph(f"Simulation Type: {analysis_type.replace('_', ' ').title()}")
        
        doc.add_heading("Simulation Summary Metrics", level=2)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Metric"
        hdr_cells[1].text = "Value"
        hdr_cells[2].text = "Unit"
        
        def add_row(m, v, u):
            row_cells = table.add_row().cells
            row_cells[0].text = str(m)
            row_cells[1].text = str(v)
            row_cells[2].text = str(u)
            
        if results.get("max_stress"):
            add_row("Max von Mises Stress", f"{results.get('max_stress'):,.2f}", "MPa")
        if results.get("max_displacement"):
            add_row("Max Displacement", f"{results.get('max_displacement'):,.4f}", "mm")
        if results.get("min_safety_factor"):
            add_row("Min Safety Factor", f"{results.get('min_safety_factor'):,.2f}", "-")
        if results.get("max_temperature"):
            add_row("Max Temperature", f"{results.get('max_temperature'):,.1f}", "°C")
            
        doc.add_paragraph("\nReport generated automatically via CAD-CAE dashboard.")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        return f"CAD-CAE Simulation Analysis Report\nProject: {project_name}\nType: {analysis_type}\nError generating DOCX: {e}".encode()


# Enforce login screen in main content area when not logged in
# if st.session_state.get("user") is None:
#     st.title("Automated CAD-to-CAE Platform")
#     st.info("👋 Please log in or sign up from the sidebar to access your projects and files.")


# ── Sidebar Configurations ────────────────────────────────────────────────────
with st.sidebar:
    st.image("https://img.icons8.com/nolan/96/artificial-intelligence.png", width=64)
    st.title("CAD-CAE Dashboard")
    st.write("---")

    # User Authentication UI Commented Out
    # if st.session_state.user is None:
    #     st.subheader("🔐 User Workspace Login")
    #     auth_action = st.radio("Choose Action", ["Login", "Sign Up"], horizontal=True)
    #     username = st.text_input("Username", key="auth_username")
    #     password = st.text_input("Password", type="password", key="auth_password")
    # 
    #     if st.button("Submit"):
    #         if not username or not password:
    #             st.error("Please enter both username and password")
    #         else:
    #             if auth_action == "Login":
    #                 if st.session_state.op_mode == "Standalone":
    #                     uid = check_standalone_login(username, password)
    #                     if uid:
    #                         st.session_state.user = username
    #                         st.session_state.user_id = uid
    #                         st.success(f"Logged in as {username}")
    #                         st.rerun()
    #                     else:
    #                         st.error("Invalid username or password")
    #                 else:
    #                     tok, uid = check_api_login(username, password)
    #                     if tok:
    #                         st.session_state.user = username
    #                         st.session_state.user_id = uid
    #                         st.session_state.token = tok
    #                         st.success(f"Logged in as {username}")
    #                         st.rerun()
    #                     else:
    #                         st.error("Invalid username or password")
    #             else:  # Sign Up
    #                 if st.session_state.op_mode == "Standalone":
    #                     uid = create_standalone_user(username, password)
    #                     if uid == "exists":
    #                         st.error("Username already exists")
    #                     elif uid:
    #                         st.session_state.user = username
    #                         st.session_state.user_id = uid
    #                         st.success(f"Registered and logged in as {username}")
    #                         st.rerun()
    #                     else:
    #                         st.error("Failed to create user")
    #                 else:
    #                     res = create_api_user(username, password)
    #                     if res == "exists":
    #                         st.error("Username already exists")
    #                     elif res:
    #                         # Log them in automatically
    #                         tok, uid = check_api_login(username, password)
    #                         if tok:
    #                             st.session_state.user = username
    #                             st.session_state.user_id = uid
    #                             st.session_state.token = tok
    #                             st.success(f"Registered and logged in as {username}")
    #                             st.rerun()
    #                     else:
    #                         st.error("Failed to create user")
    #     st.write("---")
    #     st.subheader("🔑 Or Sign-in via OAuth")
    #     oauth_provider = st.selectbox("Select Provider", ["Google", "GitHub", "Microsoft"])
    #     
    #     provider_key = oauth_provider.lower()
    #     callback_url = f"{st.session_state.api_url}/auth/callback"
    #     
    #     import httpx
    #     try:
    #         res = httpx.get(f"{st.session_state.api_url}/auth/{provider_key}", params={"redirect_uri": callback_url})
    #         if res.status_code == 200:
    #             auth_url = res.json().get("url")
    #             st.markdown(f"[🔗 Complete Sign-In with {oauth_provider}]({auth_url})")
    #     except Exception:
    #         pass
    # 
    #     if st.button(f"Mock {oauth_provider} Callback (Instant Login)"):
    #         try:
    #             cb_res = httpx.post(f"{st.session_state.api_url}/auth/callback", json={
    #                 "provider": provider_key,
    #                 "code": f"mock_{provider_key}_code_123"
    #             })
    #             if cb_res.status_code == 200:
    #                 data = cb_res.json()
    #                 st.session_state.user = data["username"]
    #                 st.session_state.user_id = data["user_id"]
    #                 st.session_state.token = data["token"]
    #                 st.success(f"Logged in via OAuth as {data['username']}")
    #                 st.rerun()
    #             else:
    #                 st.error("OAuth callback exchange failed")
    #         except Exception as e:
    #             st.error(f"OAuth connection failed: {e}")
    #     
    #     # Stop executing the rest of the dashboard
    #     st.stop()
    # 
    # else:
    #     st.write(f"👤 Logged in as: **{st.session_state.user}**")
    #     if st.button("Log Out"):
    #         st.session_state.user = None
    #         st.session_state.user_id = None
    #         st.session_state.token = None
    #         st.session_state.active_project_id = None
    #         st.session_state.active_job_id = None
    #         st.session_state.chat_messages = []
    #         st.rerun()
    #     st.write("---")

    if st.button("🆕 Start New Project"):
        st.session_state.active_project_id = None
        st.session_state.active_job_id = None
        st.session_state.chat_messages = []
        st.session_state.active_tab = "📐 Geometry Analysis"
        st.rerun()

    st.write("---")
    st.subheader("⚙️ Chatbot Settings")
    
    has_key = False
    masked_key = ""
    if st.session_state.op_mode == "API Client" and st.session_state.token:
        headers = {"Authorization": f"Bearer {st.session_state.token}"}
        try:
            res = httpx.get(f"{st.session_state.api_url}/auth/api-key", headers=headers)
            if res.status_code == 200:
                data = res.json()
                has_key = data.get("has_key", False)
                masked_key = data.get("masked_key", "")
        except Exception:
            pass

    if has_key:
        st.success(f"Registered Key: `{masked_key}`")
        if st.button("🗑️ Delete API Key"):
            if st.session_state.op_mode == "API Client" and st.session_state.token:
                headers = {"Authorization": f"Bearer {st.session_state.token}"}
                try:
                    res = httpx.delete(f"{st.session_state.api_url}/auth/api-key", headers=headers)
                    if res.status_code == 200:
                        st.success("API key deleted!")
                        st.rerun()
                except Exception as e:
                    st.error(f"Failed to delete key: {e}")
    else:
        new_key = st.text_input("Enter Gemini API Key", type="password")
        if st.button("💾 Save API Key"):
            if new_key:
                if st.session_state.op_mode == "API Client" and st.session_state.token:
                    headers = {"Authorization": f"Bearer {st.session_state.token}"}
                    try:
                        res = httpx.post(f"{st.session_state.api_url}/auth/api-key", json={"gemini_api_key": new_key}, headers=headers)
                        if res.status_code == 200:
                            st.success("API key saved securely!")
                            st.rerun()
                        else:
                            st.error("Failed to save API key.")
                    except Exception as e:
                        st.error(f"Failed to save key: {e}")
            else:
                st.warning("Please enter a key first.")


    st.subheader("🎨 Customize Interface")
    theme_choice = st.selectbox(
        "Theme Mode",
        options=["Dark Mode", "Light Mode"],
        index=0 if st.session_state.theme == "Dark" else 1
    )
    new_theme = "Dark" if theme_choice == "Dark Mode" else "Light"
    if new_theme != st.session_state.theme:
        st.session_state.theme = new_theme
        st.rerun()



# ── Fetch Active Project Details ─────────────────────────────────────────────
active_project = None
if st.session_state.active_project_id:
    if st.session_state.op_mode == "Standalone" and STANDALONE_AVAILABLE:
        active_project = get_standalone_project(st.session_state.active_project_id, st.session_state.user_id)
    elif st.session_state.op_mode == "API Client":
        active_project = api_request("GET", f"/projects/{st.session_state.active_project_id}")
        if active_project:
            # Also get geometry features
            geo = api_request("GET", f"/projects/{st.session_state.active_project_id}/geometry")
            sug = api_request("GET", f"/projects/{st.session_state.active_project_id}/suggestions")
            active_project["geometry"] = geo or {}
            active_project["geometry"]["suggestions"] = sug.get("suggestions", []) if sug else []


# ── Main Header ───────────────────────────────────────────────────────────────
if active_project:
    st.title(f"Project: {active_project['name']}")
    st.caption(f"CAD File: {active_project.get('cad_filename', 'N/A')} ({active_project.get('cad_format', '').upper()})")
else:
    st.title("Automated CAD-to-CAE Platform")
    st.write("Upload a CAD model to automatically run feature detection, obtain simulation suggestives, run multiphysics solvers, and check results.")

# ── Set Up Main Tabs ──────────────────────────────────────────────────────────
nav_cols = st.columns(4)
tab_titles = [
    "📐 Geometry Analysis",
    "💻 Simulation Setup",
    "📊 3D Viewer & Results",
    "🤖 AI Assistant Chat"
]

for idx, title in enumerate(tab_titles):
    with nav_cols[idx]:
        is_allowed = True
        err_msg = ""
        if title == "💻 Simulation Setup" and not active_project:
            is_allowed = False
            err_msg = "Please upload a CAD model first."
        elif title == "📊 3D Viewer & Results":
            jobs = []
            if active_project:
                if st.session_state.op_mode == "Standalone" and STANDALONE_AVAILABLE:
                    jobs = get_standalone_jobs(st.session_state.active_project_id)
                elif st.session_state.op_mode == "API Client":
                    job_res = api_request("GET", f"/projects/{st.session_state.active_project_id}/jobs")
                    if job_res: jobs = job_res
            completed_jobs = [j for j in jobs if j["status"] == "completed"]
            if not completed_jobs:
                is_allowed = False
                err_msg = "Please run a simulation successfully first."

        if st.button(title, use_container_width=True, type="primary" if st.session_state.active_tab == title else "secondary"):
            if is_allowed:
                st.session_state.active_tab = title
                st.rerun()
            else:
                st.error(err_msg)
st.write("---")

# 🛠 Tab 1: Geometry Feature Extraction & Suggestions
if st.session_state.active_tab == "📐 Geometry Analysis":
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📤 Upload New CAD Model")
        with st.form("cad_upload_form", clear_on_submit=True):
            uploaded_file = st.file_uploader(
                "Upload a CAD file", 
                type=None
            )
            proj_name_input = st.text_input("Project Name (Optional)", placeholder="My Analysis Project")
            submit_upload = st.form_submit_button("Analyze Geometry")

            if submit_upload and uploaded_file is not None:
                file_bytes = uploaded_file.read()
                filename = uploaded_file.name
                
                with st.spinner("Parsing and extracting geometric features..."):
                    if st.session_state.op_mode == "Standalone" and STANDALONE_AVAILABLE:
                        try:
                            project_id = local_upload_cad(file_bytes, filename, proj_name_input, st.session_state.user_id)
                            st.session_state.active_project_id = project_id
                            st.success("CAD parsed and loaded successfully in Standalone mode!")
                            st.session_state.active_tab = "💻 Simulation Setup"
                            st.rerun()
                        except Exception as e:
                            st.error(f"Failed to process CAD file: {e}")
                    elif st.session_state.op_mode == "API Client":
                        files = {"file": (filename, file_bytes)}
                        data = {"project_name": proj_name_input}
                        res = api_request("POST", "/upload", files=files, data=data)
                        if res:
                            st.session_state.active_project_id = res["project_id"]
                            st.success("CAD uploaded successfully to API server!")
                            st.session_state.active_tab = "💻 Simulation Setup"
                            st.rerun()

        # Display geometry attributes if loaded
        if active_project and "geometry" in active_project and active_project["geometry"]:
            geo = active_project["geometry"]
            st.write("---")
            st.subheader("📐 Detected Geometry Characteristics")
            
            # Grid display
            g1, g2, g3 = st.columns(3)
            with g1:
                st.markdown(f'<div class="metric-card"><div class="metric-title">Volume</div><div class="metric-value">{geo.get("volume", 0.0):,.1f} <span style="font-size:12px;color:#8ba3c7">mm³</span></div></div>', unsafe_allow_html=True)
            with g2:
                st.markdown(f'<div class="metric-card"><div class="metric-title">Surface Area</div><div class="metric-value">{geo.get("surface_area", 0.0):,.1f} <span style="font-size:12px;color:#8ba3c7">mm²</span></div></div>', unsafe_allow_html=True)
            with g3:
                # bounding box
                bbox = geo.get("bounding_box", {})
                if not bbox and "bbox_x" in geo:
                    bbox = {"x": geo["bbox_x"], "y": geo["bbox_y"], "z": geo["bbox_z"]}
                dx = bbox.get("x", 0.0)
                dy = bbox.get("y", 0.0)
                dz = bbox.get("z", 0.0)
                st.markdown(f'<div class="metric-card"><div class="metric-title">Bounding Box</div><div class="metric-value" style="font-size:16px; margin-top:12px">{dx:.1f} × {dy:.1f} × {dz:.1f} <span style="font-size:11px;color:#8ba3c7">mm</span></div></div>', unsafe_allow_html=True)

            # Feature flags
            st.markdown("##### Geometric Classification Tags:")
            tags = []
            if geo.get("is_thin_walled"): tags.append("Thin-Walled / Shell candidate")
            if geo.get("is_sheet_metal"): tags.append("Sheet Metal structure")
            if geo.get("has_holes"): tags.append("Bolt holes / Curvature detected")
            if geo.get("has_fillets"): tags.append("Fillet transitions detected")
            if geo.get("has_internal_cavity"): tags.append("Internal Cavity flow region")
            if geo.get("has_symmetry"): tags.append("Symmetrical properties")
            
            if tags:
                for tag in tags:
                    st.info(f"🔹 {tag}")
            else:
                st.write("Standard solid mechanical component classification.")

    with col2:
        st.subheader("🔍 3D Model Tessellation")
        # Load mesh vertices and faces
        vertices, faces = None, None
        if active_project:
            with st.spinner("Downloading 3D geometry..."):
                if st.session_state.op_mode == "Standalone" and STANDALONE_AVAILABLE:
                    try:
                        filepath = Path(active_project["cad_file_path"])
                        geom_data = run_async(parse_cad_file(filepath))
                        if geom_data.vertices is not None:
                            vertices = geom_data.vertices.tolist()
                            faces = geom_data.faces.tolist()
                    except Exception as e:
                        st.warning(f"Could not construct local mesh: {e}")
                elif st.session_state.op_mode == "API Client":
                    mesh_res = api_request("GET", f"/projects/{st.session_state.active_project_id}/mesh")
                    if mesh_res:
                        vertices = mesh_res.get("vertices")
                        faces = mesh_res.get("faces")

            if vertices and faces:
                fig = get_plotly_mesh(vertices, faces)
                st.plotly_chart(fig, use_container_width=True, key="geometry_mesh_viewer")
            else:
                st.info("Upload an STL/OBJ or select a project to display the 3D viewer.")

    # Suggestions row
    if active_project and "geometry" in active_project and active_project["geometry"]:
        st.write("---")
        st.subheader("💡 Recommended CAE Simulations")
        sugs = active_project["geometry"].get("suggestions", [])
        
        if sugs:
            cols = st.columns(min(len(sugs), 3))
            for idx, s in enumerate(sugs):
                col_idx = idx % len(cols)
                with cols[col_idx]:
                    conf = s.get("confidence", 0.8)
                    badge_class = "badge-high" if conf >= 0.8 else ("badge-medium" if conf >= 0.5 else "badge-low")
                    st.markdown(f"""
                    <div class="suggestion-card">
                        <div class="suggestion-title">⚙️ {s.get("title", s.get("analysis_type"))}</div>
                        <div style="margin-top:6px; margin-bottom:8px">
                            <span class="badge {badge_class}">Confidence: {conf*100:.0f}%</span>
                            <span class="badge" style="background:rgba(41,121,255,0.15);color:#2979ff;border:1px solid rgba(41,121,255,0.3)">Category: {s.get("category", "FEA")}</span>
                        </div>
                        <p style="font-size:12px; line-height:1.4">{s.get("rationale")}</p>
                    </div>
                    """, unsafe_allow_html=True)
        else:
            st.info("No recommendations found for this geometry.")


# 💻 Tab 2: Solver Parameter Configurations & Job submission
elif st.session_state.active_tab == "💻 Simulation Setup":
    if not active_project:
        st.warning("Please select or upload a project first.")
    else:
        st.subheader("⚙️ Setup Simulation Run")
        
        # Select analysis types
        sug_types = [s.get("analysis_type") for s in active_project["geometry"].get("suggestions", [])]
        all_types = [e.value for e in AnalysisType]
        # sort so suggestions appear first
        sorted_types = sorted(all_types, key=lambda x: x not in sug_types)
        
        selected_types = st.multiselect(
            "Select Simulation Type(s)",
            options=sorted_types,
            default=[sorted_types[0]] if sorted_types else [],
            format_func=lambda x: x.replace("_", " ").title()
        )

        st.markdown("### ⚙️ Mesh Settings")
        col_m1, col_m2 = st.columns(2)
        with col_m1:
            global_element_size = st.number_input("Global Element Size (mm)", min_value=0.5, max_value=50.0, value=5.0, step=0.5)
        with col_m2:
            element_order = st.selectbox("Element Order", options=[1, 2], format_func=lambda x: "Linear (C3D8)" if x == 1 else "Quadratic (C3D20R) — Recommended")
        
        # Load material options
        materials = []
        if st.session_state.op_mode == "Standalone" and STANDALONE_AVAILABLE:
            materials = get_standalone_materials()
        elif st.session_state.op_mode == "API Client":
            mat_res = api_request("GET", "/materials")
            if mat_res: materials = mat_res
            
        mat_options = {m["name"]: m for m in materials}
        
        col_mat, col_bc = st.columns(2)
        
        with col_mat:
            st.subheader("🧱 Material Assignment")
            selected_mat_name = st.selectbox("Select Material from Library", options=list(mat_options.keys()) + ["+ Add Custom Material"])
            
            if selected_mat_name == "+ Add Custom Material":
                st.markdown("**Create Custom Material**")
                c_name = st.text_input("Material Name", placeholder="e.g. Titanium Grade 5")
                c_youngs = st.number_input("Young's Modulus (GPa)", value=200.0)
                c_poissons = st.number_input("Poisson's Ratio", value=0.26, step=0.01)
                c_density = st.number_input("Density (kg/m³)", value=7850.0)
                c_yield = st.number_input("Yield Strength (MPa)", value=250.0)
                c_cond = st.number_input("Thermal Conductivity (W/m·K)", value=50.0)
                c_spec = st.number_input("Specific Heat (J/kg·K)", value=490.0)
                
                assigned_mat = {
                    "name": c_name or "Custom",
                    "category": "Metal",
                    "youngs_modulus": c_youngs,
                    "poissons_ratio": c_poissons,
                    "density": c_density,
                    "yield_strength": c_yield,
                    "thermal_conductivity": c_cond,
                    "specific_heat": c_spec
                }

                if st.button("➕ Save Custom Material"):
                    if not c_name.strip():
                        st.error("Material name is required.")
                    elif c_youngs <= 0 or c_poissons < 0 or c_poissons >= 0.5 or c_density <= 0 or c_yield <= 0 or c_cond <= 0 or c_spec <= 0:
                        st.error("Please enter valid positive numbers for material properties.")
                    else:
                        if st.session_state.op_mode == "Standalone" and STANDALONE_AVAILABLE:
                            async def save_custom():
                                async with get_db_context() as db:
                                    from app.models.db_models import Material as DbMaterial
                                    mat_obj = DbMaterial(
                                        name=c_name,
                                        category="Metal",
                                        youngs_modulus=c_youngs,
                                        poissons_ratio=c_poissons,
                                        density=c_density,
                                        yield_strength=c_yield,
                                        thermal_conductivity=c_cond,
                                        specific_heat=c_spec,
                                        is_custom=True,
                                        user_id=st.session_state.user_id
                                    )
                                    db.add(mat_obj)
                            run_async(save_custom())
                            st.success("Custom material saved in Standalone mode!")
                            st.rerun()
                        elif st.session_state.op_mode == "API Client":
                            res = api_request("POST", "/materials", json=assigned_mat)
                            if res:
                                st.success("Custom material saved to API server!")
                                st.rerun()
            else:
                # Show assigned properties
                if selected_mat_name:
                    assigned_mat = mat_options[selected_mat_name]
                    st.markdown(f"""
                    - **Young's Modulus**: {assigned_mat.get('youngs_modulus', 'N/A')} GPa
                    - **Poisson's Ratio**: {assigned_mat.get('poissons_ratio', 'N/A')}
                    - **Density**: {assigned_mat.get('density', 'N/A')} kg/m³
                    - **Yield Strength**: {assigned_mat.get('yield_strength', 'N/A')} MPa
                    - **Thermal Conductivity**: {assigned_mat.get('thermal_conductivity', 'N/A')} W/(m·K)
                    """)

        with col_bc:
            st.subheader("⚓ Boundary Conditions")
            
            bcs = []
            fatigue_bcs = {}
            for atype in selected_types:
                if "thermal" in atype:
                    st.markdown(f"**Thermal Settings ({atype.replace('_', ' ').title()}):**")
                    temp_inlet = st.number_input("Source Temperature (°C)", value=100.0, key=f"temp_inlet_{atype}")
                    temp_ambient = st.number_input("Ambient/Convection Temp (°C)", value=20.0, key=f"temp_ambient_{atype}")
                    convection_h = st.number_input("Convective Heat Transfer Coeff H (W/m²K)", value=15.0, key=f"convection_h_{atype}")
                    bcs.extend([
                        {"type": "temperature", "value": temp_inlet, "location": "boundary_inlet"},
                        {"type": "convection", "ambient_temp": temp_ambient, "h": convection_h, "location": "boundary_ambient"}
                    ])
                elif "cfd" in atype:
                    st.markdown(f"**CFD Settings ({atype.replace('_', ' ').title()}):**")
                    inlet_vel = st.number_input("Inlet Flow Velocity (m/s)", value=2.5, key=f"inlet_vel_{atype}")
                    outlet_p = st.number_input("Outlet Relative Pressure (Pa)", value=0.0, key=f"outlet_p_{atype}")
                    bcs.extend([
                        {"type": "velocity_inlet", "value": inlet_vel, "location": "inlet"},
                        {"type": "pressure_outlet", "value": outlet_p, "location": "outlet"}
                    ])
                elif atype == "fatigue":
                    st.markdown(f"**Fatigue Settings ({atype.replace('_', ' ').title()}):**")
                    stress_amp = st.number_input("Cyclic Stress Amplitude (MPa)", value=100.0, key=f"stress_amp_{atype}")
                    stress_ratio = st.number_input("Stress Ratio R (Min/Max stress)", value=-1.0, key=f"stress_ratio_{atype}")
                    mean_correction = st.selectbox("Mean Stress Correction Method", options=["goodman", "soderberg", "gerber"], key=f"mean_correction_{atype}")
                    fatigue_bcs = {
                        "stress_amplitude": stress_amp,
                        "stress_ratio": stress_ratio,
                        "mean_stress_correction": mean_correction
                    }
                else:  # Structural
                    st.markdown(f"**Structural Settings ({atype.replace('_', ' ').title()}):**")
                    force_z = st.number_input("Load Force in Z-direction (N)", value=-1000.0, key=f"force_z_{atype}")
                    force_y = st.number_input("Load Force in Y-direction (N)", value=0.0, key=f"force_y_{atype}")
                    force_x = st.number_input("Load Force in X-direction (N)", value=0.0, key=f"force_x_{atype}")
                    bcs.extend([
                        {"type": "fixed", "location": "fixed_support"},
                        {"type": "force", "fx": force_x, "fy": force_y, "fz": force_z, "location": "load_point"}
                    ])
                
            # Pack parameter payload
            parameters = {
                "material": assigned_mat if selected_mat_name else {},
                "boundary_conditions": bcs,
                "mesh_settings": {
                    "global_element_size": global_element_size,
                    "refinement_factor": 1.0,
                    "element_order": element_order
                }
            }
            if fatigue_bcs:
                parameters.update(fatigue_bcs)
                
        # Solver dispatch triggers
        st.write("---")
        if st.button("🚀 Solve Multiphysics Model"):
            job_id = str(uuid.uuid4())
            
            with st.spinner("Dispatching solver run..."):
                if st.session_state.op_mode == "Standalone" and STANDALONE_AVAILABLE:
                    # Create job record
                    async def save_job():
                        async with get_db_context() as db:
                            job = Job(
                                id=job_id,
                                project_id=st.session_state.active_project_id,
                                analysis_types=selected_types,
                                status=JobStatus.PENDING,
                                parameters=parameters
                            )
                            db.add(job)
                    run_async(save_job())
                    
                    # Run inline solver in a separate thread to keep streamlit loop reactive
                    import threading
                    mesh_filepath = active_project.get("cad_file_path")
                    async def run_solver():
                        await run_analysis_inline(job_id, selected_types, parameters, mesh_filepath)
                    
                    thread = threading.Thread(target=run_async, args=(run_solver(),))
                    thread.start()
                    
                    # Show progress updates in real-time
                    prog_bar = st.progress(0.0, text="Initializing runner...")
                    status_lbl = st.empty()
                    
                    while thread.is_alive():
                        job_info = get_standalone_job(job_id)
                        if job_info:
                            pct = job_info.get("progress_percent", 0)
                            msg = job_info.get("progress_message", "Running...")
                            prog_bar.progress(pct / 100.0, text=f"{msg} ({pct}%)")
                            if job_info.get("status") in ("completed", "failed", "cancelled"):
                                break
                        time.sleep(0.5)
                        
                    # Final check
                    job_info = get_standalone_job(job_id)
                    if job_info and job_info.get("status") == "completed":
                        st.session_state.active_job_id = job_id
                        st.success("Simulation finished successfully!")
                        st.session_state.active_tab = "📊 3D Viewer & Results"
                        st.rerun()
                    else:
                        st.error(f"Simulation failed. Error: {job_info.get('error_message') if job_info else 'Unknown'}")
                        
                elif st.session_state.op_mode == "API Client":
                    res = api_request("POST", f"/projects/{st.session_state.active_project_id}/jobs", json={
                        "analysis_types": selected_types,
                        "parameters": parameters
                    })
                    if res:
                        job_id = res["id"]
                        prog_bar = st.progress(0.0, text="Dispatched...")
                        
                        # Poll REST progress endpoint
                        for _ in range(200): # 100 seconds timeout
                            time.sleep(0.5)
                            job_info = api_request("GET", f"/jobs/{job_id}")
                            if job_info:
                                pct = job_info.get("progress_percent", 0)
                                msg = job_info.get("progress_message", "Executing...")
                                prog_bar.progress(pct / 100.0, text=f"{msg} ({pct}%)")
                                if job_info.get("status") in ("completed", "failed", "cancelled"):
                                    break
                        
                        if job_info and job_info.get("status") == "completed":
                            st.session_state.active_job_id = job_id
                            st.success("Simulation completed on server!")
                            st.session_state.active_tab = "📊 3D Viewer & Results"
                            st.rerun()
                        else:
                            st.error("Simulation run timed out or failed on the API server.")


# 📊 Tab 3: Simulation Results, Contour plots and PDF reports
elif st.session_state.active_tab == "📊 3D Viewer & Results":
    # Let user select from the history of completed jobs for this project
    jobs = []
    if active_project:
        if st.session_state.op_mode == "Standalone" and STANDALONE_AVAILABLE:
            jobs = get_standalone_jobs(st.session_state.active_project_id)
        elif st.session_state.op_mode == "API Client":
            job_res = api_request("GET", f"/projects/{st.session_state.active_project_id}/jobs")
            if job_res: jobs = job_res
            
    completed_jobs = [j for j in jobs if j["status"] == "completed"]
    
    if not completed_jobs:
        st.info("No completed simulation runs found. Configure and run a simulation in Tab 2 first.")
    else:
        job_options = {}
        for j in completed_jobs:
            types_list = j.get("analysis_types") or [j.get("analysis_type", "static_structural")]
            lbl = f"{', '.join(types_list).replace('_', ' ').title()} - {j['id'][:8]}"
            job_options[lbl] = j["id"]
        selected_job_lbl = st.selectbox("Select Simulation Results to view", options=list(job_options.keys()))
        active_job_id = job_options[selected_job_lbl]
        
        # Load results payload
        active_job = None
        if st.session_state.op_mode == "Standalone" and STANDALONE_AVAILABLE:
            active_job = get_standalone_job(active_job_id)
        elif st.session_state.op_mode == "API Client":
            active_job = api_request("GET", f"/jobs/{active_job_id}")
            if active_job:
                # also get full results schema
                res_payload = api_request("GET", f"/jobs/{active_job_id}/results")
                active_job["results"] = res_payload or {}
                
        if active_job and "results" in active_job and active_job["results"]:
            res = active_job["results"]
            
            # If multiple analysis types, show selector
            active_types = active_job.get("analysis_types") or [active_job.get("analysis_type", "static_structural")]
            if len(active_types) > 1:
                result_type = st.radio("Select Analysis Type Result to View", options=active_types, format_func=lambda x: x.replace("_", " ").title())
            else:
                result_type = active_types[0] if active_types else "static_structural"
                
            st.write("---")
            st.subheader(f"📊 Key Simulation Scalar Metrics ({result_type.replace('_', ' ').title()})")
            
            # Extract nested/merged summary or fallback to top level
            res_data_dict = res.get("result_data", {}).get("results", {}).get(result_type, {})
            summary_dict = res.get("result_data", {}).get("summaries", {}).get(result_type, {})
            if not summary_dict and not res_data_dict:
                # fallback to legacy
                summary_dict = res
                res_data_dict = res.get("result_data", {})

            m1, m2, m3, m4 = st.columns(4)
            with m1:
                val = summary_dict.get("max_stress") or res_data_dict.get("max_stress")
                if val is not None: st.markdown(f'<div class="metric-card"><div class="metric-title">Max von Mises Stress</div><div class="metric-value">{val:,.2f} <span style="font-size:12px;color:#8ba3c7">MPa</span></div></div>', unsafe_allow_html=True)
            with m2:
                val = summary_dict.get("max_displacement") or res_data_dict.get("max_displacement")
                if val is not None: st.markdown(f'<div class="metric-card"><div class="metric-title">Max Displacement</div><div class="metric-value">{val:,.4f} <span style="font-size:12px;color:#8ba3c7">mm</span></div></div>', unsafe_allow_html=True)
            with m3:
                val = summary_dict.get("min_safety_factor")
                if val is not None:
                    sf_color = "#ff5252" if val < 1.2 else ("#ffab40" if val < 2.0 else "#00e676")
                    st.markdown(f'<div class="metric-card"><div class="metric-title">Min Safety Factor</div><div class="metric-value" style="color:{sf_color}">{val:,.2f}</div></div>', unsafe_allow_html=True)
            with m4:
                val = summary_dict.get("max_temperature") or res_data_dict.get("max_temperature")
                if val is not None: st.markdown(f'<div class="metric-card"><div class="metric-title">Max Temperature</div><div class="metric-value">{val:,.1f} <span style="font-size:12px;color:#8ba3c7">°C</span></div></div>', unsafe_allow_html=True)
                
            # Modal natural frequencies
            freqs = summary_dict.get("natural_frequencies") or res_data_dict.get("natural_frequencies_hz")
            if freqs:
                st.info(f"🎵 Detected Natural Frequencies (First {len(freqs)} modes): " + ", ".join([f"{f:.1f} Hz" for f in freqs]))
                
            # Fatigue cycles
            cycles = summary_dict.get("fatigue_life_cycles") or res_data_dict.get("fatigue_life_cycles")
            if cycles:
                st.warning(f"⏳ Min Fatigue Lifetime: {cycles:,.0f} cycles")
                
            # Buckling factors
            bf = summary_dict.get("buckling_factors") or res_data_dict.get("buckling_factors")
            if bf:
                st.info(f"⚖️ Buckling Eigenvalue Load Factors: " + ", ".join([f"{f:.2f}" for f in bf]))

            st.write("---")
            st.subheader("👁️ 3D Interactive Results Contour Map")
            
            # Choose color field
            color_field_choices = ["stress", "displacement", "safety_factor"]
            if "thermal" in result_type:
                color_field_choices = ["temperature"]
            elif "cfd" in result_type:
                color_field_choices = ["velocity"]
                
            field_choice = st.selectbox("Select Result Field Contour", options=color_field_choices)
            
            # Draw contour plot
            vertices, faces = None, None
            if st.session_state.op_mode == "Standalone" and STANDALONE_AVAILABLE:
                try:
                    filepath = Path(active_project["cad_file_path"])
                    geom_data = run_async(parse_cad_file(filepath))
                    if geom_data.vertices is not None:
                        vertices = geom_data.vertices.tolist()
                        faces = geom_data.faces.tolist()
                except Exception as e:
                    pass
            elif st.session_state.op_mode == "API Client":
                mesh_res = api_request("GET", f"/projects/{st.session_state.active_project_id}/mesh")
                if mesh_res:
                    vertices = mesh_res.get("vertices")
                    faces = mesh_res.get("faces")

            slice_axis = "None"
            slice_val = None
            if vertices and faces and "cfd" in result_type:
                st.markdown("**✂️ CFD Slicing Cutaway**")
                slice_axis = st.selectbox("Select Slice Axis", options=["None", "X-Axis", "Y-Axis", "Z-Axis"])
                if slice_axis != "None":
                    axis_idx = 0 if slice_axis == "X-Axis" else (1 if slice_axis == "Y-Axis" else 2)
                    coords = np.array(vertices)
                    min_coord = float(np.min(coords[:, axis_idx]))
                    max_coord = float(np.max(coords[:, axis_idx]))
                    slice_val = st.slider("Slice Location", min_value=min_coord, max_value=max_coord, value=(min_coord + max_coord) / 2)

            if vertices and faces:
                viz_data = {**summary_dict, **res_data_dict}
                fig = get_plotly_mesh(vertices, faces, result_type, viz_data, field_choice, slice_axis=slice_axis, slice_val=slice_val)
                st.plotly_chart(fig, use_container_width=True, key="results_mesh_viewer")
            else:
                st.error("Failed to load CAD mesh vertices for results post-processing.")
                
            st.write("---")
            st.subheader("📄 Engineering Report Generation")
            col_pdf, col_docx = st.columns(2)
            
            project_name = active_project.get("name", "Unknown")
            viz_data = {**summary_dict, **res_data_dict}
            
            # Generate report bytes using backend generator
            pdf_data = b""
            docx_data = b""
            if st.session_state.op_mode == "Standalone" and STANDALONE_AVAILABLE:
                async def _gen():
                    async with get_db_context() as db:
                        from app.results.report_generator import generate_reports_for_job
                        pdf_p, docx_p = await generate_reports_for_job(active_job_id, db)
                        p_pdf = pdf_p.read_bytes() if pdf_p and pdf_p.exists() else b""
                        p_docx = docx_p.read_bytes() if docx_p and docx_p.exists() else b""
                        return p_pdf, p_docx
                pdf_data, docx_data = run_async(_gen())
            else:
                headers = {}
                if st.session_state.token:
                    headers["Authorization"] = f"Bearer {st.session_state.token}"
                try:
                    r_pdf = httpx.get(f"{st.session_state.api_url}/api/jobs/{active_job_id}/results/pdf", headers=headers, timeout=30.0)
                    if r_pdf.status_code == 200:
                        pdf_data = r_pdf.content
                    r_docx = httpx.get(f"{st.session_state.api_url}/api/jobs/{active_job_id}/results/docx", headers=headers, timeout=30.0)
                    if r_docx.status_code == 200:
                        docx_data = r_docx.content
                except Exception as e:
                    pass
            
            with col_pdf:
                st.download_button(
                    label="📥 Download PDF Analysis Report",
                    data=pdf_data,
                    file_name=f"{project_name.lower().replace(' ', '_')}_{result_type}_report.pdf",
                    mime="application/pdf"
                )
            with col_docx:
                st.download_button(
                    label="📥 Download Word DOCX Analysis Report",
                    data=docx_data,
                    file_name=f"{project_name.lower().replace(' ', '_')}_{result_type}_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )


# 🤖 Tab 4: AI chatbot Grounded in CAE project context
elif st.session_state.active_tab == "🤖 AI Assistant Chat":
    if not active_project:
        st.warning("Please select or upload a project first.")
    else:
        st.subheader("💬 AI Engineering Assistant")
        st.caption("Ask questions about boundary conditions, stress levels, material alternatives, design optimizations, or help interpreting simulation plots.")
        
        # Display chat history
        for msg in st.session_state.chat_messages:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])
                
        # Handle chat input
        user_input = st.chat_input("Ask a question about the current analysis...")
        
        if user_input:
            # Render user message
            with st.chat_message("user"):
                st.write(user_input)
            st.session_state.chat_messages.append({"role": "user", "content": user_input})
            
            # Prepare context dictionary matching backend _build_chat_context
            active_job_context = {}
            if st.session_state.active_job_id:
                if st.session_state.op_mode == "Standalone" and STANDALONE_AVAILABLE:
                    job_data = get_standalone_job(st.session_state.active_job_id)
                    if job_data and "results" in job_data:
                        r = job_data["results"]
                        active_job_context = {
                            "job": {
                                "type": job_data["analysis_type"],
                                "status": job_data["status"],
                            },
                            "results": {
                                "max_stress_mpa": r.get("max_stress"),
                                "max_displacement_mm": r.get("max_displacement"),
                                "safety_factor": r.get("min_safety_factor"),
                            }
                        }
                elif st.session_state.op_mode == "API Client":
                    job_data = api_request("GET", f"/jobs/{st.session_state.active_job_id}")
                    if job_data:
                        r_data = api_request("GET", f"/jobs/{st.session_state.active_job_id}/results")
                        r = r_data or {}
                        active_job_context = {
                            "job": {
                                "type": job_data["analysis_type"],
                                "status": job_data["status"],
                            },
                            "results": {
                                "max_stress_mpa": r.get("max_stress"),
                                "max_displacement_mm": r.get("max_displacement"),
                                "safety_factor": r.get("min_safety_factor"),
                            }
                        }

            # Build full context
            context = {
                "project_name": active_project["name"],
                "cad_format": active_project.get("cad_format", ""),
                "geometry": {
                    "volume_mm3": active_project["geometry"].get("volume"),
                    "surface_area_mm2": active_project["geometry"].get("surface_area"),
                    "is_thin_walled": active_project["geometry"].get("is_thin_walled"),
                    "has_holes": active_project["geometry"].get("has_holes"),
                },
                "history": st.session_state.chat_messages
            }
            context.update(active_job_context)
            
            # Streaming assistant response
            with st.chat_message("assistant"):
                token_placeholder = st.empty()
                full_response = ""
                
                # Standalone mode: stream directly if Claude is configured
                if st.session_state.op_mode == "Standalone" and STANDALONE_AVAILABLE:
                    if not settings.anthropic_api_key:
                        full_response = (
                            "I'm the CAE Analysis Assistant. The Claude API key is not configured in the `.env` file yet.\n\n"
                            "Based on your geometry:\n"
                            f"- The volume is **{context['geometry']['volume_mm3'] or 0:,.1f} mm³**\n"
                            f"- The surface area is **{context['geometry']['surface_area_mm2'] or 0:,.1f} mm²**\n"
                            f"- Thin-Walled classification: **{context['geometry']['is_thin_walled']}**\n\n"
                            "Please configure `ANTHROPIC_API_KEY` to enable active interactive reasoning."
                        )
                        token_placeholder.write(full_response)
                    else:
                        async def stream_from_claude():
                            response_text = ""
                            async for token in stream_chat_response(user_input, context):
                                response_text += token
                                token_placeholder.write(response_text)
                            return response_text
                        full_response = run_async(stream_from_claude())
                        
                elif st.session_state.op_mode == "API Client":
                    # Call API chat router and parse streaming response
                    try:
                        response_text = ""
                        for token in api_chat_stream(st.session_state.api_url, st.session_state.active_project_id, st.session_state.active_job_id, user_input):
                            response_text += token
                            token_placeholder.write(response_text)
                        full_response = response_text
                    except Exception as e:
                        full_response = f"Could not stream response from API: {e}"
                        token_placeholder.write(full_response)
                        
                st.session_state.chat_messages.append({"role": "assistant", "content": full_response})
