"""
Results parser and report generator.
Generates fully-detailed PDF and Word reports summarizing materials, boundary conditions, mesh, and results.
"""
from __future__ import annotations
import json
from pathlib import Path
from datetime import datetime, timezone
import structlog
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models import db_models as M

logger = structlog.get_logger(__name__)


async def generate_reports_for_job(job_id: str, db: AsyncSession) -> tuple[Path | None, Path | None]:
    """Generates both PDF and DOCX reports for a job and updates the DB paths."""
    job = (await db.execute(select(M.Job).where(M.Job.id == job_id))).scalar_one_or_none()
    if not job:
        return None, None

    project = (await db.execute(select(M.Project).where(M.Project.id == job.project_id))).scalar_one_or_none()
    geo = (await db.execute(select(M.GeometryFeatures).where(M.GeometryFeatures.project_id == job.project_id))).scalar_one_or_none()
    res = (await db.execute(select(M.AnalysisResult).where(M.AnalysisResult.job_id == job_id))).scalar_one_or_none()

    if not res:
        return None, None

    # Determine paths
    storage_dir = Path(settings.local_storage_path) / "reports"
    storage_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = storage_dir / f"{job_id}_report.pdf"
    docx_path = storage_dir / f"{job_id}_report.docx"

    try:
        # Build PDF
        _build_pdf(pdf_path, job, project, geo, res)
        res.report_pdf_path = str(pdf_path)
    except Exception as e:
        logger.error("Failed to generate PDF report", job_id=job_id, error=str(e))
        pdf_path = None

    try:
        # Build DOCX
        _build_docx(docx_path, job, project, geo, res)
        res.report_docx_path = str(docx_path)
    except Exception as e:
        logger.error("Failed to generate DOCX report", job_id=job_id, error=str(e))
        docx_path = None

    await db.commit()
    return pdf_path, docx_path


def _build_pdf(output_path: Path, job: M.Job, project: M.Project | None, geo: M.GeometryFeatures | None, res: M.AnalysisResult):
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    doc = SimpleDocTemplate(str(output_path), pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    story = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        name='TitleStyle',
        parent=styles['Heading1'],
        textColor=colors.HexColor('#002d62'),
        spaceAfter=15
    )
    section_style = ParagraphStyle(
        name='SectionStyle',
        parent=styles['Heading2'],
        textColor=colors.HexColor('#004b93'),
        spaceBefore=12,
        spaceAfter=6
    )

    # Title & Metadata
    story.append(Paragraph("CAD-CAE Simulation Analysis Report", title_style))
    story.append(Paragraph(f"<b>Job ID:</b> {job.id}", styles['Normal']))
    story.append(Paragraph(f"<b>Project Name:</b> {project.name if project else 'N/A'}", styles['Normal']))
    story.append(Paragraph(f"<b>Generated At:</b> {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}", styles['Normal']))
    story.append(Spacer(1, 15))

    # Section 1: Analysis Inputs
    story.append(Paragraph("1. Analysis Inputs", section_style))
    story.append(Spacer(1, 4))

    # Geometry / Mesh inputs
    mesh_size = "N/A"
    element_order = "N/A"
    materials_list = []
    bcs_list = []

    if job.parameters:
        for atype, params in job.parameters.items():
            if not isinstance(params, dict):
                continue
            # Material
            mat = params.get("material")
            if mat and isinstance(mat, dict):
                mat_name = mat.get("name", "Unknown")
                mat_props = []
                if mat.get("youngs_modulus"): mat_props.append(f"E: {mat.get('youngs_modulus')} MPa")
                if mat.get("poissons_ratio"): mat_props.append(f"Nu: {mat.get('poissons_ratio')}")
                if mat.get("density"): mat_props.append(f"Density: {mat.get('density')} kg/mm³")
                mat_desc = f"{mat_name} ({', '.join(mat_props)})"
                if mat_desc not in materials_list:
                    materials_list.append(mat_desc)

            # Mesh Settings
            mesh_set = params.get("mesh_settings")
            if mesh_set and isinstance(mesh_set, dict):
                mesh_size = f"{mesh_set.get('global_element_size', '5.0')} mm"
                element_order = f"Quadratic (Order {mesh_set.get('element_order', 2)})"

            # BCs
            bcs = params.get("boundary_conditions")
            if bcs and isinstance(bcs, list):
                for bc in bcs:
                    if not isinstance(bc, dict):
                        continue
                    bc_type = bc.get("type", "Unknown").replace("_", " ").title()
                    bc_faces = f"Faces: {bc.get('face_ids', [])}"
                    bc_vals = []
                    for k, v in bc.items():
                        if k not in ("type", "face_ids", "description") and isinstance(v, (int, float)):
                            bc_vals.append(f"{k}: {v}")
                    bc_desc = f"{bc_type} on {bc_faces}"
                    if bc_vals:
                        bc_desc += f" ({', '.join(bc_vals)})"
                    if bc_desc not in bcs_list:
                        bcs_list.append(bc_desc)

    inputs_data = [
        ["Input Parameter", "Value / Details"],
        ["Analysis Types Run", ", ".join(job.analysis_types).replace("_", " ").title()],
        ["Material Name(s)", "\n".join(materials_list) if materials_list else "N/A"],
        ["Boundary Conditions", "\n".join(bcs_list) if bcs_list else "N/A"],
        ["Target Global Mesh Size", mesh_size],
        ["Element Order / Type", element_order],
        ["Element Count", f"{geo.element_count:,}" if geo and geo.element_count else "N/A"],
        ["Node Count", f"{geo.node_count:,}" if geo and geo.node_count else "N/A"]
    ]

    t_inputs = Table(inputs_data, colWidths=[180, 320])
    t_inputs.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#004b93')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 6),
        ('TOPPADDING', (0,0), (-1,0), 6),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#d0d5dd')),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')])
    ]))
    story.append(t_inputs)
    story.append(Spacer(1, 15))

    # Section 2: Results Summary
    story.append(Paragraph("2. Simulation Scalar Results", section_style))
    story.append(Spacer(1, 4))

    results_data = [["Metric", "Value", "Unit"]]
    if res.max_stress is not None:
        results_data.append(["Max von Mises Stress", f"{res.max_stress:,.2f}", "MPa"])
    if res.min_stress is not None:
        results_data.append(["Min von Mises Stress", f"{res.min_stress:,.2f}", "MPa"])
    if res.max_displacement is not None:
        results_data.append(["Max Displacement", f"{res.max_displacement:,.4f}", "mm"])
    if res.min_safety_factor is not None:
        results_data.append(["Min Safety Factor", f"{res.min_safety_factor:,.2f}", "-"])
    if res.max_temperature is not None:
        results_data.append(["Max Temperature", f"{res.max_temperature:,.1f}", "°C"])
    if res.natural_frequencies:
        results_data.append(["Natural Frequencies (Modes 1-3)", ", ".join([f"{f:.1f} Hz" for f in res.natural_frequencies[:3]]), "Hz"])
    if res.buckling_factors:
        results_data.append(["Buckling Load Multipliers", ", ".join([f"{f:.2f}x" for f in res.buckling_factors[:3]]), "x"])
    if res.fatigue_life_cycles is not None:
        results_data.append(["Min Fatigue Life", f"{res.fatigue_life_cycles:,.0f}", "cycles"])

    t_res = Table(results_data, colWidths=[200, 200, 100])
    t_res.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#2979ff')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 6),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#d0d5dd')),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')])
    ]))
    story.append(t_res)
    story.append(Spacer(1, 20))
    story.append(Paragraph("<i>Disclaimer: This report is a decision-support helper. Final engineering decisions require approval by a licensed professional engineer.</i>", styles['Italic']))

    doc.build(story)


def _build_pdf(output_path: Path, job: M.Job, project: M.Project | None, geo: M.GeometryFeatures | None, res: M.AnalysisResult):
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    doc = SimpleDocTemplate(str(output_path), pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    story = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        name='TitleStyle',
        parent=styles['Heading1'],
        textColor=colors.HexColor('#002d62'),
        spaceAfter=15
    )
    section_style = ParagraphStyle(
        name='SectionStyle',
        parent=styles['Heading2'],
        textColor=colors.HexColor('#004b93'),
        spaceBefore=12,
        spaceAfter=6
    )

    # Title & Metadata
    story.append(Paragraph("CAD-CAE Simulation Analysis Report", title_style))
    story.append(Paragraph(f"<b>Job ID:</b> {job.id}", styles['Normal']))
    story.append(Paragraph(f"<b>Project Name:</b> {project.name if project else 'N/A'}", styles['Normal']))
    story.append(Paragraph(f"<b>Generated At:</b> {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}", styles['Normal']))
    story.append(Spacer(1, 15))

    # Section 1: Analysis Inputs
    story.append(Paragraph("1. Analysis Inputs", section_style))
    story.append(Spacer(1, 4))

    # Geometry / Mesh inputs
    mesh_size = "N/A"
    element_order = "N/A"
    materials_list = []
    bcs_list = []

    # Check if parameters are nested by analysis type or flat
    params_list = []
    if job.parameters:
        if "material" in job.parameters or "boundary_conditions" in job.parameters or "mesh_settings" in job.parameters:
            params_list = [job.parameters]
        else:
            params_list = [p for p in job.parameters.values() if isinstance(p, dict)]

    for params in params_list:
        # Material
        mat = params.get("material")
        if mat and isinstance(mat, dict):
            mat_name = mat.get("name", "Unknown")
            mat_props = []
            if mat.get("youngs_modulus"): mat_props.append(f"E: {mat.get('youngs_modulus')} GPa")
            if mat.get("poissons_ratio"): mat_props.append(f"Nu: {mat.get('poissons_ratio')}")
            if mat.get("density"): mat_props.append(f"Density: {mat.get('density')} kg/m³")
            if mat.get("yield_strength"): mat_props.append(f"Yield: {mat.get('yield_strength')} MPa")
            mat_desc = f"{mat_name} ({', '.join(mat_props)})"
            if mat_desc not in materials_list:
                materials_list.append(mat_desc)

        # Mesh Settings
        mesh_set = params.get("mesh_settings")
        if mesh_set and isinstance(mesh_set, dict):
            mesh_size = f"{mesh_set.get('global_element_size', '5.0')} mm"
            element_order = f"Quadratic (Order {mesh_set.get('element_order', 2)})" if mesh_set.get('element_order') == 2 else f"Linear (Order {mesh_set.get('element_order', 1)})"

        # BCs
        bcs = params.get("boundary_conditions")
        if bcs and isinstance(bcs, list):
            for bc in bcs:
                if not isinstance(bc, dict):
                    continue
                bc_type = bc.get("type", "Unknown").replace("_", " ").title()
                bc_faces = f"Faces: {bc.get('face_ids', [])}"
                if bc.get("description"):
                    bc_faces = f"{bc.get('description')} ({bc_faces})"
                bc_vals = []
                for k, v in bc.items():
                    if k not in ("type", "face_ids", "description") and isinstance(v, (int, float, str)):
                        bc_vals.append(f"{k}: {v}")
                bc_desc = f"{bc_type} on {bc_faces}"
                if bc_vals:
                    bc_desc += f" ({', '.join(bc_vals)})"
                if bc_desc not in bcs_list:
                    bcs_list.append(bc_desc)

    inputs_data = [
        ["Input Parameter", "Value / Details"],
        ["Analysis Types Run", ", ".join(job.analysis_types).replace("_", " ").title()],
        ["Material Properties", "\n".join(materials_list) if materials_list else "N/A"],
        ["Boundary Conditions", "\n".join(bcs_list) if bcs_list else "N/A"],
        ["Target Global Mesh Size", mesh_size],
        ["Element Order / Type", element_order],
        ["Element Count", f"{geo.element_count:,}" if geo and geo.element_count is not None else "N/A"],
        ["Node Count", f"{geo.node_count:,}" if geo and geo.node_count is not None else "N/A"]
    ]

    t_inputs = Table(inputs_data, colWidths=[180, 320])
    t_inputs.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#004b93')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 6),
        ('TOPPADDING', (0,0), (-1,0), 6),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#d0d5dd')),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')])
    ]))
    story.append(t_inputs)
    story.append(Spacer(1, 15))

    # Mesh Quality Parameters section
    story.append(Paragraph("2. Mesh Quality Parameters", section_style))
    story.append(Spacer(1, 4))
    
    quality_data = None
    if geo and geo.raw_features:
        quality_data = geo.raw_features.get("mesh_quality_metrics")
        
    if quality_data:
        ar = quality_data.get("aspect_ratio", {})
        skew = quality_data.get("skewness", {})
        jac = quality_data.get("jacobian_ratio", {})
        ang = quality_data.get("angles", {})
        flagged = quality_data.get("flagged_counts", {})
        
        quality_data_rows = [
            ["Metric", "Value / Details"],
            ["Aspect Ratio (Mean / Max)", f"{ar.get('mean', 1.0):.2f} / {ar.get('max', 1.0):.2f}"],
            ["Skewness (Mean / Max)", f"{skew.get('mean', 0.0):.2f} / {skew.get('max', 0.0):.2f}"],
            ["Jacobian Ratio (Mean / Max)", f"{jac.get('mean', 1.0):.2f} / {jac.get('max', 1.0):.2f}"],
            ["Min / Max Angle", f"{ang.get('min_angle_min', 60.0):.1f}° / {ang.get('max_angle_max', 60.0):.1f}°"],
            ["Flagged Elements (Aspect Ratio / Skew)", f"{flagged.get('aspect_ratio', 0)} / {flagged.get('skewness', 0)}"],
            ["Total Flagged Elements", f"{flagged.get('total_flagged', 0)} ({quality_data.get('overall_quality_score', 1.0)*100:.1f}% acceptable)"]
        ]
    else:
        quality_data_rows = [
            ["Metric", "Value / Details"],
            ["Mesh Quality Metrics", "N/A - Run Mesh Refinement first"]
        ]
        
    t_quality = Table(quality_data_rows, colWidths=[200, 300])
    t_quality.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#6b7280')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 6),
        ('TOPPADDING', (0,0), (-1,0), 6),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#d0d5dd')),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')])
    ]))
    story.append(t_quality)
    story.append(Spacer(1, 15))

    # Section 3: Results Summary
    story.append(Paragraph("3. Simulation Scalar Results", section_style))
    story.append(Spacer(1, 4))

    results_data = [["Metric", "Value", "Unit"]]
    if res.max_stress is not None:
        results_data.append(["Max von Mises Stress", f"{res.max_stress:,.2f}", "MPa"])
    if res.min_stress is not None:
        results_data.append(["Min von Mises Stress", f"{res.min_stress:,.2f}", "MPa"])
    if res.max_displacement is not None:
        results_data.append(["Max Displacement", f"{res.max_displacement:,.4f}", "mm"])
    if res.min_safety_factor is not None:
        results_data.append(["Min Safety Factor", f"{res.min_safety_factor:,.2f}", "-"])
    if res.max_temperature is not None:
        results_data.append(["Max Temperature", f"{res.max_temperature:,.1f}", "°C"])
    if res.natural_frequencies:
        results_data.append(["Natural Frequencies (Modes 1-3)", ", ".join([f"{f:.1f} Hz" for f in res.natural_frequencies[:3]]), "Hz"])
    if res.buckling_factors:
        results_data.append(["Buckling Load Multipliers", ", ".join([f"{f:.2f}x" for f in res.buckling_factors[:3]]), "x"])
    if res.fatigue_life_cycles is not None:
        results_data.append(["Min Fatigue Life", f"{res.fatigue_life_cycles:,.0f}", "cycles"])

    t_res = Table(results_data, colWidths=[200, 200, 100])
    t_res.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#2979ff')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0,0), (-1,0), 6),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#d0d5dd')),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')])
    ]))
    story.append(t_res)
    story.append(Spacer(1, 15))

    # Section 4: Visual Contour Snapshot
    image_path = Path(settings.local_storage_path) / "reports" / f"{job.id}_contour.png"
    if image_path.exists():
        story.append(Paragraph("4. Results Contour Visualization", section_style))
        story.append(Spacer(1, 4))
        try:
            story.append(Image(str(image_path), width=400, height=250))
            story.append(Spacer(1, 15))
        except Exception as e:
            logger.error("Failed to embed image in PDF report", job_id=job.id, error=str(e))

    story.append(Paragraph("<i>Disclaimer: This report is a decision-support helper. Final engineering decisions require approval by a licensed professional engineer.</i>", styles['Italic']))

    doc.build(story)


def _build_docx(output_path: Path, job: M.Job, project: M.Project | None, geo: M.GeometryFeatures | None, res: M.AnalysisResult):
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("CAD-CAE Simulation Analysis Report", level=0)

    # Metadata
    doc.add_paragraph(f"Job ID: {job.id}")
    doc.add_paragraph(f"Project Name: {project.name if project else 'N/A'}")
    doc.add_paragraph(f"Generated At: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")

    # Section 1: Analysis Inputs
    doc.add_heading("1. Analysis Inputs", level=1)
    
    mesh_size = "N/A"
    element_order = "N/A"
    materials_list = []
    bcs_list = []

    # Check if parameters are nested or flat
    params_list = []
    if job.parameters:
        if "material" in job.parameters or "boundary_conditions" in job.parameters or "mesh_settings" in job.parameters:
            params_list = [job.parameters]
        else:
            params_list = [p for p in job.parameters.values() if isinstance(p, dict)]

    for params in params_list:
        mat = params.get("material")
        if mat and isinstance(mat, dict):
            mat_name = mat.get("name", "Unknown")
            mat_props = []
            if mat.get("youngs_modulus"): mat_props.append(f"E: {mat.get('youngs_modulus')} GPa")
            if mat.get("poissons_ratio"): mat_props.append(f"Nu: {mat.get('poissons_ratio')}")
            if mat.get("density"): mat_props.append(f"Density: {mat.get('density')} kg/m³")
            if mat.get("yield_strength"): mat_props.append(f"Yield: {mat.get('yield_strength')} MPa")
            mat_desc = f"{mat_name} ({', '.join(mat_props)})"
            if mat_desc not in materials_list:
                materials_list.append(mat_desc)

        mesh_set = params.get("mesh_settings")
        if mesh_set and isinstance(mesh_set, dict):
            mesh_size = f"{mesh_set.get('global_element_size', '5.0')} mm"
            element_order = f"Quadratic (Order {mesh_set.get('element_order', 2)})" if mesh_set.get('element_order') == 2 else f"Linear (Order {mesh_set.get('element_order', 1)})"

        bcs = params.get("boundary_conditions")
        if bcs and isinstance(bcs, list):
            for bc in bcs:
                if not isinstance(bc, dict):
                    continue
                bc_type = bc.get("type", "Unknown").replace("_", " ").title()
                bc_faces = f"Faces: {bc.get('face_ids', [])}"
                if bc.get("description"):
                    bc_faces = f"{bc.get('description')} ({bc_faces})"
                bc_vals = []
                for k, v in bc.items():
                    if k not in ("type", "face_ids", "description") and isinstance(v, (int, float, str)):
                        bc_vals.append(f"{k}: {v}")
                bc_desc = f"{bc_type} on {bc_faces}"
                if bc_vals:
                    bc_desc += f" ({', '.join(bc_vals)})"
                if bc_desc not in bcs_list:
                    bcs_list.append(bc_desc)

    table_inputs = doc.add_table(rows=1, cols=2)
    table_inputs.style = 'Table Grid'
    hdr_cells = table_inputs.rows[0].cells
    hdr_cells[0].text = "Input Parameter"
    hdr_cells[1].text = "Value / Details"

    inputs = [
        ("Analysis Types Run", ", ".join(job.analysis_types).replace("_", " ").title()),
        ("Material Properties", "\n".join(materials_list) if materials_list else "N/A"),
        ("Boundary Conditions", "\n".join(bcs_list) if bcs_list else "N/A"),
        ("Target Global Mesh Size", mesh_size),
        ("Element Order / Type", element_order),
        ("Element Count", f"{geo.element_count:,}" if geo and geo.element_count is not None else "N/A"),
        ("Node Count", f"{geo.node_count:,}" if geo and geo.node_count is not None else "N/A")
    ]

    for param, val in inputs:
        row_cells = table_inputs.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = val

    # Mesh Quality Parameters section
    doc.add_heading("2. Mesh Quality Parameters", level=1)
    
    quality_data = None
    if geo and geo.raw_features:
        quality_data = geo.raw_features.get("mesh_quality_metrics")
        
    table_quality = doc.add_table(rows=1, cols=2)
    table_quality.style = 'Table Grid'
    hdr_cells = table_quality.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "Value / Details"
    
    if quality_data:
        ar = quality_data.get("aspect_ratio", {})
        skew = quality_data.get("skewness", {})
        jac = quality_data.get("jacobian_ratio", {})
        ang = quality_data.get("angles", {})
        flagged = quality_data.get("flagged_counts", {})
        
        quality_rows = [
            ("Aspect Ratio (Mean / Max)", f"{ar.get('mean', 1.0):.2f} / {ar.get('max', 1.0):.2f}"),
            ("Skewness (Mean / Max)", f"{skew.get('mean', 0.0):.2f} / {skew.get('max', 0.0):.2f}"),
            ("Jacobian Ratio (Mean / Max)", f"{jac.get('mean', 1.0):.2f} / {jac.get('max', 1.0):.2f}"),
            ("Min / Max Angle", f"{ang.get('min_angle_min', 60.0):.1f}° / {ang.get('max_angle_max', 60.0):.1f}°"),
            ("Flagged Elements (Aspect Ratio / Skew)", f"{flagged.get('aspect_ratio', 0)} / {flagged.get('skewness', 0)}"),
            ("Total Flagged Elements", f"{flagged.get('total_flagged', 0)} ({quality_data.get('overall_quality_score', 1.0)*100:.1f}% acceptable)")
        ]
    else:
        quality_rows = [
            ("Mesh Quality Metrics", "N/A - Run Mesh Refinement first")
        ]
        
    for param, val in quality_rows:
        row_cells = table_quality.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = val

    # Section 3: Results
    doc.add_heading("3. Simulation Scalar Results", level=1)
    
    table_res = doc.add_table(rows=1, cols=3)
    table_res.style = 'Table Grid'
    hdr_cells = table_res.rows[0].cells
    hdr_cells[0].text = "Metric"
    hdr_cells[1].text = "Value"
    hdr_cells[2].text = "Unit"

    def add_res_row(m, v, u):
        row_cells = table_res.add_row().cells
        row_cells[0].text = str(m)
        row_cells[1].text = str(v)
        row_cells[2].text = str(u)

    if res.max_stress is not None:
        add_res_row("Max von Mises Stress", f"{res.max_stress:,.2f}", "MPa")
    if res.min_stress is not None:
        add_res_row("Min von Mises Stress", f"{res.min_stress:,.2f}", "MPa")
    if res.max_displacement is not None:
        add_res_row("Max Displacement", f"{res.max_displacement:,.4f}", "mm")
    if res.min_safety_factor is not None:
        add_res_row("Min Safety Factor", f"{res.min_safety_factor:,.2f}", "-")
    if res.max_temperature is not None:
        add_res_row("Max Temperature", f"{res.max_temperature:,.1f}", "°C")
    if res.natural_frequencies:
        add_res_row("Natural Frequencies (Modes 1-3)", ", ".join([f"{f:.1f} Hz" for f in res.natural_frequencies[:3]]), "Hz")
    if res.buckling_factors:
        add_res_row("Buckling Load Multipliers", ", ".join([f"{f:.2f}x" for f in res.buckling_factors[:3]]), "x")
    if res.fatigue_life_cycles is not None:
        add_res_row("Min Fatigue Life", f"{res.fatigue_life_cycles:,.0f}", "cycles")

    # Section 4: Visual Contour Snapshot
    image_path = Path(settings.local_storage_path) / "reports" / f"{job.id}_contour.png"
    if image_path.exists():
        doc.add_heading("4. Results Contour Visualization", level=1)
        try:
            doc.add_picture(str(image_path), width=Inches(5.5))
        except Exception as e:
            doc.add_paragraph(f"Error loading visualization image: {e}")

    doc.add_paragraph("\nDisclaimer: This report is a decision-support helper. Final engineering decisions require approval by a licensed professional engineer.")
    
    doc.save(str(output_path))
